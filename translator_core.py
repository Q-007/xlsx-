import pandas as pd
from deep_translator import GoogleTranslator
import time
import threading
import os
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

class TranslatorBase:
    def __init__(self, source_lang='auto', target_lang='en'):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.translator = GoogleTranslator(source=source_lang, target=target_lang)
        self.stop_requested = False

    def translate_text(self, text):
        if not isinstance(text, str) or not text.strip():
            return text
        try:
            # Simple retry mechanism
            for attempt in range(3):
                try:
                    return self.translator.translate(text)
                except Exception as e:
                    if attempt == 2:
                        print(f"Failed to translate '{text}': {e}")
                        return text
                    time.sleep(1)
        except Exception:
            return text
            
    def stop(self):
        self.stop_requested = True

class ExcelTranslator(TranslatorBase):
    def process_file(self, input_path, output_path, progress_callback=None, status_callback=None):
        try:
            xls = pd.ExcelFile(input_path)
            # Use openpyxl engine for writing to support .xlsx
            writer = pd.ExcelWriter(output_path, engine='openpyxl')
            
            total_sheets = len(xls.sheet_names)
            
            for sheet_idx, sheet_name in enumerate(xls.sheet_names):
                if self.stop_requested:
                    break
                    
                df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
                
                total_cells = df.size
                processed_cells = 0
                
                if status_callback:
                    status_callback(f"正在处理工作表: {sheet_name} ({sheet_idx + 1}/{total_sheets})")

                def translate_cell(cell):
                    nonlocal processed_cells
                    if self.stop_requested:
                        return cell
                    
                    processed_cells += 1
                    
                    if processed_cells % 10 == 0 or processed_cells == total_cells:
                        if progress_callback:
                            sheet_progress = processed_cells / total_cells
                            total_progress = (sheet_idx + sheet_progress) / total_sheets * 100
                            progress_callback(total_progress)

                    return self.translate_text(cell)

                translated_df = df.map(translate_cell)
                translated_df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
            
            writer.close()
            
            if self.stop_requested:
                if status_callback: status_callback("翻译已被用户停止。")
                return False
            
            if status_callback: status_callback("Excel 翻译成功完成！")
            if progress_callback: progress_callback(100)
            return True

        except Exception as e:
            if status_callback: status_callback(f"错误: {str(e)}")
            return False

class WordTranslator(TranslatorBase):
    def process_file(self, input_path, output_path, progress_callback=None, status_callback=None):
        if Document is None:
            if status_callback: status_callback("错误: 未安装 python-docx 库。")
            return False
            
        try:
            doc = Document(input_path)
            
            # Count total elements (paragraphs + table cells) for progress
            total_elements = len(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    total_elements += len(row.cells)
            
            processed_count = 0
            
            if status_callback: status_callback("正在处理 Word 文档...")

            # Helper to update progress
            def update_progress():
                nonlocal processed_count
                processed_count += 1
                if progress_callback and (processed_count % 5 == 0 or processed_count == total_elements):
                    progress_callback(min(processed_count / total_elements * 100, 99))

            # Translate paragraphs
            for paragraph in doc.paragraphs:
                if self.stop_requested: break
                
                if paragraph.text.strip():
                    translated = self.translate_text(paragraph.text)
                    paragraph.text = translated
                update_progress()

            # Translate tables
            for table in doc.tables:
                if self.stop_requested: break
                for row in table.rows:
                    if self.stop_requested: break
                    for cell in row.cells:
                        if self.stop_requested: break
                        if cell.text.strip():
                            translated = self.translate_text(cell.text)
                            cell.text = translated
                        update_progress()
            
            if self.stop_requested:
                if status_callback: status_callback("翻译已被用户停止。")
                return False

            doc.save(output_path)
            
            if status_callback: status_callback("Word 翻译成功完成！")
            if progress_callback: progress_callback(100)
            return True

        except Exception as e:
            if status_callback: status_callback(f"错误: {str(e)}")
            return False

class PDFTranslator(TranslatorBase):
    def process_file(self, input_path, output_path, progress_callback=None, status_callback=None):
        if Converter is None:
            if status_callback: status_callback("错误: 未安装 pdf2docx 库。")
            return False
            
        try:
            # Step 1: Convert PDF to temporary Word file
            if status_callback: status_callback("正在将 PDF 转换为 Word 格式（以保留排版）...")
            if progress_callback: progress_callback(10) # Initial progress
            
            temp_docx = input_path + ".temp.docx"
            
            # Suppress pdf2docx output
            import sys
            import io
            original_stdout = sys.stdout
            sys.stdout = io.StringIO()
            
            try:
                cv = Converter(input_path)
                cv.convert(temp_docx, start=0, end=None)
                cv.close()
            except Exception as e:
                sys.stdout = original_stdout
                raise e
            finally:
                sys.stdout = original_stdout
            
            if self.stop_requested:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                if status_callback: status_callback("翻译已被用户停止。")
                return False
                
            if status_callback: status_callback("PDF 转换完成，开始翻译内容...")
            
            # Step 2: Use WordTranslator to translate the temporary file
            # Output will be a .docx file (user can export to PDF if needed, but docx is better for editing)
            word_translator = WordTranslator(self.source_lang, self.target_lang)
            
            # Proxy callbacks to adjust progress range (10% -> 100%)
            def proxy_progress(val):
                if progress_callback:
                    # Map 0-100 to 10-100
                    real_progress = 10 + (val * 0.9)
                    progress_callback(real_progress)
            
            success = word_translator.process_file(
                temp_docx, 
                output_path, 
                progress_callback=proxy_progress, 
                status_callback=lambda msg: status_callback(f"{msg}") if "Word" not in msg else None # Filter internal messages
            )
            
            # Cleanup temp file
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
            if success:
                if status_callback: status_callback("PDF 翻译完成！(已保存为可编辑的 Word 文档)")
                if progress_callback: progress_callback(100)
                return True
            else:
                return False

        except Exception as e:
            if status_callback: status_callback(f"错误: {str(e)}")
            # Cleanup temp file in case of error
            try:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
            except:
                pass
            return False
